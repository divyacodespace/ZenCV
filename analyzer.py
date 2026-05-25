"""Resume analysis engine: extracts text, scores ATS-readiness, and surfaces gaps."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Iterable

import pdfplumber
from docx import Document


SECTION_KEYWORDS = {
    "contact": ["email", "phone", "linkedin", "github", "@"],
    "summary": ["summary", "objective", "profile", "about me"],
    "experience": ["experience", "employment", "work history", "professional"],
    "education": ["education", "academic", "university", "college", "bachelor", "master", "b.tech", "m.tech", "degree"],
    "skills": ["skills", "technologies", "technical", "tools"],
    "projects": ["projects", "portfolio"],
    "certifications": ["certification", "certificate", "licensed"],
    "achievements": ["achievement", "award", "honor"],
}

COMMON_SKILLS = {
    "Programming": [
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
        "ruby", "php", "kotlin", "swift", "scala", "r",
    ],
    "Web": [
        "html", "css", "react", "angular", "vue", "node.js", "express", "django",
        "flask", "fastapi", "spring", "next.js", "tailwind", "bootstrap",
    ],
    "Data & ML": [
        "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
        "matplotlib", "seaborn", "nlp", "machine learning", "deep learning",
        "data analysis", "statistics",
    ],
    "Databases": [
        "sql", "mysql", "postgresql", "mongodb", "sqlite", "redis", "oracle",
        "dynamodb", "cassandra",
    ],
    "Cloud & DevOps": [
        "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "git", "github",
        "ci/cd", "terraform", "linux",
    ],
    "Soft Skills": [
        "communication", "leadership", "teamwork", "problem solving",
        "collaboration", "adaptability", "time management",
    ],
}

ACTION_VERBS = {
    "led", "built", "designed", "developed", "implemented", "created", "managed",
    "improved", "optimized", "increased", "reduced", "launched", "delivered",
    "engineered", "automated", "architected", "spearheaded", "achieved",
    "deployed", "migrated", "refactored", "analyzed", "researched",
}

ATS_UNFRIENDLY_PATTERNS = [
    (r"[│┃┆┊┋║]", "Vertical bar / box-drawing characters from tables"),
    (r"[•◦▪▫■□●○◆◇★☆]", "Decorative bullet glyphs (use simple - or *)"),
    (r"\t{2,}", "Multiple tabs (often from columns/tables)"),
]


@dataclass
class AnalysisResult:
    filename: str
    text: str
    word_count: int
    ats_score: int
    sections_found: list[str]
    sections_missing: list[str]
    skills_found: dict[str, list[str]]
    skills_missing: dict[str, list[str]]
    keyword_gaps: list[str]
    formatting_issues: list[str]
    strengths: list[str]
    weaknesses: list[str]
    suggestions: list[str]
    breakdown: dict[str, dict] = field(default_factory=dict)


def extract_text(file_storage) -> str:
    """Read text out of an uploaded PDF or DOCX file."""
    filename = (file_storage.filename or "").lower()
    data = file_storage.read()
    file_storage.seek(0)

    if filename.endswith(".pdf"):
        return _extract_pdf(data)
    if filename.endswith(".docx"):
        return _extract_docx(data)
    raise ValueError("Unsupported file type. Upload a .pdf or .docx file.")


def _extract_pdf(data: bytes) -> str:
    chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            chunks.append(page_text)
    return "\n".join(chunks).strip()


def _extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def _find_sections(text_lower: str) -> tuple[list[str], list[str]]:
    found, missing = [], []
    for section, keywords in SECTION_KEYWORDS.items():
        if any(k in text_lower for k in keywords):
            found.append(section)
        else:
            missing.append(section)
    return found, missing


def _find_skills(text_lower: str) -> tuple[dict[str, list[str]], dict[str, list[str]]]:
    found: dict[str, list[str]] = {}
    missing: dict[str, list[str]] = {}
    for category, skills in COMMON_SKILLS.items():
        present = [s for s in skills if s in text_lower]
        absent = [s for s in skills if s not in text_lower]
        if present:
            found[category] = present
        missing[category] = absent[:6]
    return found, missing


def _formatting_issues(text: str) -> list[str]:
    issues: list[str] = []
    for pattern, label in ATS_UNFRIENDLY_PATTERNS:
        if re.search(pattern, text):
            issues.append(label)

    if text.count("  ") > 20:
        issues.append("Excessive double-spacing — likely column layout artifacts.")
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if lines and sum(1 for ln in lines if ln.isupper()) / len(lines) > 0.25:
        issues.append("Too many ALL-CAPS lines — ATS systems prefer mixed case.")
    if "image" in text.lower() and "imagemagick" not in text.lower():
        issues.append("References to images detected — graphics may not be parsed by ATS.")
    return issues


def _quantified_achievements(text: str) -> int:
    return len(re.findall(r"\b\d+(?:[.,]\d+)?\s*(?:%|percent|k|million|years?|hours?|users?|customers?)\b",
                          text, flags=re.IGNORECASE))


def _action_verb_count(text_lower: str) -> int:
    return sum(1 for verb in ACTION_VERBS if re.search(rf"\b{re.escape(verb)}\b", text_lower))


def _contact_info(text: str) -> dict[str, bool]:
    return {
        "email": bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)),
        "phone": bool(re.search(r"(\+?\d[\d\s\-().]{7,}\d)", text)),
        "linkedin": "linkedin.com" in text.lower(),
    }


def _score(
    sections_found: list[str],
    skills_found: dict[str, list[str]],
    formatting_issues: list[str],
    contact: dict[str, bool],
    quantified: int,
    action_verbs: int,
    word_count: int,
) -> tuple[int, dict[str, dict]]:
    breakdown: dict[str, dict] = {}

    section_score = min(25, int(25 * len(sections_found) / len(SECTION_KEYWORDS)))
    breakdown["Sections"] = {"score": section_score, "max": 25,
                              "detail": f"{len(sections_found)}/{len(SECTION_KEYWORDS)} key sections present"}

    skill_total = sum(len(v) for v in skills_found.values())
    skill_score = min(25, skill_total * 2)
    breakdown["Skills coverage"] = {"score": skill_score, "max": 25,
                                     "detail": f"{skill_total} recognized skills detected"}

    formatting_score = max(0, 15 - len(formatting_issues) * 4)
    breakdown["Formatting"] = {"score": formatting_score, "max": 15,
                                "detail": f"{len(formatting_issues)} formatting issue(s) found"}

    contact_points = sum(contact.values())
    contact_score = int(10 * contact_points / 3)
    breakdown["Contact info"] = {"score": contact_score, "max": 10,
                                  "detail": f"{contact_points}/3 contact channels present"}

    quant_score = min(10, quantified * 2)
    breakdown["Quantified impact"] = {"score": quant_score, "max": 10,
                                       "detail": f"{quantified} numeric achievement(s) found"}

    verb_score = min(10, action_verbs)
    breakdown["Action verbs"] = {"score": verb_score, "max": 10,
                                  "detail": f"{action_verbs} strong action verbs used"}

    if 350 <= word_count <= 900:
        length_score = 5
        length_detail = f"{word_count} words — within recommended range"
    elif 200 <= word_count < 350 or 900 < word_count <= 1200:
        length_score = 3
        length_detail = f"{word_count} words — slightly outside ideal range"
    else:
        length_score = 1
        length_detail = f"{word_count} words — outside recommended range (350–900)"
    breakdown["Length"] = {"score": length_score, "max": 5, "detail": length_detail}

    total = sum(item["score"] for item in breakdown.values())
    return min(100, total), breakdown


def _strengths_and_weaknesses(
    sections_found: list[str],
    sections_missing: list[str],
    skills_found: dict[str, list[str]],
    formatting_issues: list[str],
    contact: dict[str, bool],
    quantified: int,
    action_verbs: int,
) -> tuple[list[str], list[str]]:
    strengths: list[str] = []
    weaknesses: list[str] = []

    if len(sections_found) >= 6:
        strengths.append("Covers most expected resume sections.")
    if skills_found:
        top = max(skills_found.items(), key=lambda kv: len(kv[1]))
        strengths.append(f"Strong coverage in {top[0]} ({len(top[1])} skills listed).")
    if quantified >= 3:
        strengths.append(f"Includes {quantified} quantified achievements — great for impact.")
    if action_verbs >= 6:
        strengths.append("Uses strong action verbs throughout.")
    if all(contact.values()):
        strengths.append("Contact information is complete (email, phone, LinkedIn).")
    if not formatting_issues:
        strengths.append("Clean, ATS-friendly formatting.")

    if "summary" in sections_missing:
        weaknesses.append("No professional summary or objective at the top.")
    if "skills" in sections_missing:
        weaknesses.append("No explicit Skills section detected.")
    if "experience" in sections_missing:
        weaknesses.append("Work experience section is missing or unclear.")
    if "projects" in sections_missing:
        weaknesses.append("No Projects section — adds credibility for technical roles.")
    if quantified < 2:
        weaknesses.append("Few quantified results — recruiters want numbers, %, and scale.")
    if action_verbs < 4:
        weaknesses.append("Weak verb usage — bullets should start with strong action verbs.")
    if not contact.get("linkedin"):
        weaknesses.append("LinkedIn profile URL is missing.")
    if not contact.get("email"):
        weaknesses.append("No email address detected.")
    if formatting_issues:
        weaknesses.extend(formatting_issues)

    return strengths, weaknesses


def _suggestions(
    skills_missing: dict[str, list[str]],
    sections_missing: list[str],
    quantified: int,
    word_count: int,
) -> tuple[list[str], list[str]]:
    suggestions: list[str] = []
    keyword_gaps: list[str] = []

    for category, items in skills_missing.items():
        if not items:
            continue
        sample = ", ".join(items[:4])
        suggestions.append(f"Consider adding relevant {category} keywords if they apply: {sample}.")
        keyword_gaps.extend(items[:4])

    if "summary" in sections_missing:
        suggestions.append("Add a 2–3 line professional summary tailored to the target role.")
    if "projects" in sections_missing:
        suggestions.append("Showcase 2–3 projects with tech stack, your role, and measurable outcomes.")
    if "certifications" in sections_missing:
        suggestions.append("List relevant certifications — they often match ATS keywords directly.")
    if quantified < 3:
        suggestions.append("Quantify achievements: 'Reduced load time by 40%', 'Led team of 5', etc.")
    if word_count < 350:
        suggestions.append("Resume seems too short — expand experience bullets with impact statements.")
    elif word_count > 1200:
        suggestions.append("Resume is dense — trim to 1–2 pages and keep only role-relevant content.")

    suggestions.append("Save and submit as PDF generated from a single-column template for best ATS parsing.")
    return suggestions, keyword_gaps


def analyze(file_storage) -> AnalysisResult:
    text = extract_text(file_storage)
    if not text or len(text.split()) < 30:
        raise ValueError("Could not extract enough text. The file may be scanned/image-based.")

    text_lower = text.lower()
    word_count = len(text.split())

    sections_found, sections_missing = _find_sections(text_lower)
    skills_found, skills_missing = _find_skills(text_lower)
    formatting_issues = _formatting_issues(text)
    contact = _contact_info(text)
    quantified = _quantified_achievements(text)
    action_verbs = _action_verb_count(text_lower)

    ats_score, breakdown = _score(
        sections_found, skills_found, formatting_issues, contact,
        quantified, action_verbs, word_count,
    )
    strengths, weaknesses = _strengths_and_weaknesses(
        sections_found, sections_missing, skills_found, formatting_issues,
        contact, quantified, action_verbs,
    )
    suggestions, keyword_gaps = _suggestions(
        skills_missing, sections_missing, quantified, word_count,
    )

    return AnalysisResult(
        filename=file_storage.filename or "resume",
        text=text,
        word_count=word_count,
        ats_score=ats_score,
        sections_found=sections_found,
        sections_missing=sections_missing,
        skills_found=skills_found,
        skills_missing=skills_missing,
        keyword_gaps=keyword_gaps,
        formatting_issues=formatting_issues,
        strengths=strengths,
        weaknesses=weaknesses,
        suggestions=suggestions,
        breakdown=breakdown,
    )
